"""Optional disclaimer table.

A template may contain ONE disclaimer table whose top-left cell starts with the
tag ``{{DISC}}``. The user chooses (a checkbox in the UI) whether the report
should include the disclaimer:

  * include = True  -> keep the table, remove just the ``{{DISC}}`` tag text.
  * include = False -> remove the whole disclaimer table.

This mirrors the tag-driven table handling already used in ``tables`` and the
block expansion in ``damage_blocks``. It runs on the already-built output doc.
"""

from docx import Document

TAG = "{{DISC}}"


def _strip_tag_in_cell(cell):
    """Remove the TAG text from a cell, collapsing across split runs.

    Word often splits a tag across several runs, so if no single run holds the
    whole tag we rebuild the first paragraph's text from the joined run text.
    """
    for para in cell.paragraphs:
        if TAG not in "".join(r.text for r in para.runs):
            continue
        # Fast path: a single run contains the whole tag.
        if any(TAG in r.text for r in para.runs):
            for run in para.runs:
                if TAG in run.text:
                    run.text = run.text.replace(TAG, "")
            continue
        # Split across runs: join into the first run, blank the rest.
        joined = "".join(r.text for r in para.runs)
        para.runs[0].text = joined.replace(TAG, "")
        for run in para.runs[1:]:
            run.text = ""


def apply_disclaimer(doc, include):
    """Apply the disclaimer choice to `doc` in place.

    Returns True if the {{DISC}} table was found (kept or removed), else False.
    """
    for table in doc.tables:
        first_cell = table.rows[0].cells[0]
        if TAG not in first_cell.text:
            continue
        if include:
            _strip_tag_in_cell(first_cell)
        else:
            table._tbl.getparent().remove(table._tbl)
        return True
    return False


def apply_in_file(path, include, progress=None, review=None, doc=None):
    """Apply the disclaimer choice. When `doc` is given, operate on it and do
    not save (caller owns the single open/save); otherwise open/save `path`."""
    log = progress or (lambda m: None)
    rev = review or (lambda m: None)

    own = doc is None
    if own:
        doc = Document(path)
    found = apply_disclaimer(doc, include)
    if own:
        doc.save(path)

    if found:
        log(f"Disclaimer: {'kept' if include else 'removed'} (tag {TAG}).")
    elif include:
        rev(f"⚠ Disclaimer requested, but the template has no {TAG} table — "
            f"nothing added.")
    else:
        log(f"Disclaimer: no {TAG} table (none requested).")
    return found

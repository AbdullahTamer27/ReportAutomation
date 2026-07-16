"""Interval table — the dynamic, transposed, block-structured depth table.

Unlike every other report table (one styled data row cloned per record), the
interval table lays intervals out ACROSS columns, three per block, with a
repeating set of labelled rows: start/end depth, one "Tubular size & weight" row
per pipe, and the channel/offset rows. The template carries a SINGLE styled
block as a prototype, marked by ``{{INTERVALS}}`` in its "Intervals" label cell.
This pass clones that block once per group of three intervals, sizing the
tubular rows to each block's pipe count, and fills the cells from the computed
interval table.

The interval data is the SAME table the RawData workbook is built from
(``build_intervals_from_xml``), so the report and the Excel can never disagree.
Two deliberate rules (per the report owner):
  * "Pipe channel response" (the first-response channels) is filled from a
    ``Channels`` sheet in the data workbook when present — one already-formatted
    string per interval, interval 1 at the top — and left blank otherwise;
  * values are written as computed — anomalies (e.g. an undetected boundary) are
    left for the user to edit rather than second-guessed here.
"""

import copy
import os

from docx import Document
from docx.table import _Row

from .tables import set_cell_text

INTERVALS_TAG = "{{INTERVALS}}"
PER_BLOCK = 3   # intervals shown per block (data columns 1..3)

# First-response channels feed the "Pipe channel response" row. When the data
# workbook has a sheet named `Channels`, each interval's value is read as an
# already-formatted string (e.g. "10-20-30-40") from a single column, one row per
# interval, interval 1 at START_ROW (matching the interval numbering: shallowest
# first). The column is a PLACEHOLDER until confirmed — change it in one place.
FIRST_RESPONSE_SHEET = "Channels"
FIRST_RESPONSE_COL = "Q"         # column holding each interval's first-response string
FIRST_RESPONSE_START_ROW = 2     # interval 1 is row 2 (row 1 is the header)

# Column-0 labels, in block order. "Tubular size & weight" repeats — one row per
# pipe. "Pipe channel response" is written blank on purpose (see module docstring).
LBL_INTERVALS = "Intervals"
LBL_START = "Start Depth(ft)"
LBL_END = "End Depth(ft)"
LBL_TUBULAR = "Tubular size & weight"
LBL_INTERP = "Interpretation Channels"
LBL_RESPONSE = "Pipe channel response"
LBL_OFFSET = "Offset"
LBL_WELL_NAME = "Well Name"


# ---------------- data ----------------
def build_interval_records(xml_path, excel_path=None):
    """Compute the interval rows, reusing the exact core the RawData workbook
    uses (so the report table and the Excel can never disagree). Returns a list
    of dicts; ``Channels``/``Offsets`` are present only when a THICKNESS sheet
    was readable in `excel_path`."""
    from well_tools.core.xml_parser import parse_wellschematic_xml
    from well_tools.core.intervals import build_intervals_from_xml

    xml_data = parse_wellschematic_xml(xml_path)
    thickness_sections = None
    if excel_path and os.path.isfile(excel_path):
        try:
            from well_tools.core.thickness import parse_thickness_sections
            thickness_sections = parse_thickness_sections(excel_path) or None
        except ValueError:
            thickness_sections = None   # no/unreadable THICKNESS — channel rows stay blank
    df = build_intervals_from_xml(xml_data, thickness_sections=thickness_sections)
    records = df.to_dict("records")

    # First-response channels (the "Pipe channel response" row), aligned to the
    # same interval order — blank when there's no Channels sheet.
    responses = _read_first_response(excel_path, len(records))
    for rec, val in zip(records, responses):
        rec["FirstResponse"] = val
    return records


def _read_first_response(excel_path, count):
    """Return `count` first-response strings from the data workbook's ``Channels``
    sheet — column ``FIRST_RESPONSE_COL``, one row per interval starting at
    ``FIRST_RESPONSE_START_ROW`` (interval 1 at the top). Any missing sheet, row,
    or cell yields "" for that slot, so the row simply stays blank."""
    blanks = [""] * count
    if not excel_path or not os.path.isfile(excel_path):
        return blanks
    try:
        from openpyxl.utils import column_index_from_string
        from . import _wbcache
        wb = _wbcache.load(excel_path, data_only=True)
        name = next((s for s in wb.sheetnames
                     if s.strip().lower() == FIRST_RESPONSE_SHEET.lower()), None)
        if not name:
            return blanks
        ws = wb[name]
        col = column_index_from_string(FIRST_RESPONSE_COL)
        out = []
        for i in range(count):
            v = ws.cell(row=FIRST_RESPONSE_START_ROW + i, column=col).value
            out.append("" if v is None else str(v).strip())
        return out
    except Exception:  # noqa: BLE001 — the channel column is best-effort
        return blanks


def _fmt_depth(v):
    """Whole numbers print without a trailing '.0'; everything else as-is."""
    if v is None:
        return ""
    try:
        f = float(v)
    except (TypeError, ValueError):
        return str(v)
    return str(int(f)) if f.is_integer() else str(f)


# ---------------- prototype location ----------------
def _find_interval_table(doc):
    for t in doc.tables:
        for r in t.rows:
            if any(INTERVALS_TAG in c.text for c in r.cells):
                return t
    return None


def _prototype_block(table):
    """Locate the styled prototype block: the rows from the ``{{INTERVALS}}`` row
    through the first "Offset" row. Returns a dict of role -> <w:tr> (with
    ``tubular`` a list and ``all`` every block row), or None."""
    rows = list(table.rows)
    start = next((i for i, r in enumerate(rows)
                  if any(INTERVALS_TAG in c.text for c in r.cells)), None)
    if start is None:
        return None

    proto = {"tubular": [], "all": []}
    for r in rows[start:]:
        label = r.cells[0].text.strip()
        tr = r._tr
        proto["all"].append(tr)
        if INTERVALS_TAG in label or label == LBL_INTERVALS:
            proto["intervals"] = tr
        elif label == LBL_START:
            proto["start"] = tr
        elif label == LBL_END:
            proto["end"] = tr
        elif label == LBL_TUBULAR:
            proto["tubular"].append(tr)
        elif label == LBL_INTERP:
            proto["interp"] = tr
        elif label == LBL_RESPONSE:
            proto["response"] = tr
        elif label == LBL_OFFSET:
            proto["offset"] = tr
            break   # a block ends at its Offset row
    needed = ("intervals", "start", "end", "interp", "response", "offset")
    if not all(k in proto for k in needed) or not proto["tubular"]:
        return None
    return proto


def _fill_row(table, tr, label, values, ncols):
    """Set a block row: column 0 = its label, columns 1.. = one value per
    interval in the chunk (missing columns cleared)."""
    cells = _Row(tr, table).cells
    set_cell_text(cells[0], label)
    for j in range(1, ncols):
        i = j - 1
        set_cell_text(cells[j], values[i] if i < len(values) else "")


# ---------------- orchestration ----------------
def place_interval_table(output_path, records, well_name=None,
                         progress=None, review=None, doc=None):
    """Fill the ``{{INTERVALS}}`` table in `output_path` from `records` (a list
    of interval dicts from :func:`build_interval_records`). Clones the prototype
    block once per group of three intervals. No-ops (and leaves the document
    untouched) when the template has no ``{{INTERVALS}}`` table — so templates
    without it are unaffected. When `doc` is given, works on it and does not save.
    Returns the number of intervals written."""
    log = progress or (lambda m: None)
    rev = review or (lambda m: None)

    own = doc is None
    if own:
        doc = Document(output_path)

    table = _find_interval_table(doc)
    if table is None:
        log("Interval table: no {{INTERVALS}} table in the template — skipped.")
        if own:
            doc.save(output_path)
        return 0

    # Fill the Well Name meta cell; RIG Name is deliberately left as-is.
    if well_name:
        for r in table.rows:
            if r.cells[0].text.strip() == LBL_WELL_NAME:
                set_cell_text(r.cells[1], str(well_name))
                break

    proto = _prototype_block(table)
    if proto is None:
        rev("⚠ Interval table: prototype block incomplete — table left as-is.")
        if own:
            doc.save(output_path)
        return 0

    ncols = len(table.columns)
    proto_first = proto["intervals"]
    proto_tubular = proto["tubular"][0]
    records = list(records)
    chunks = [records[i:i + PER_BLOCK] for i in range(0, len(records), PER_BLOCK)]

    def clone_before(src_tr):
        new_tr = copy.deepcopy(src_tr)
        proto_first.addprevious(new_tr)   # keeps blocks in order, above the prototype
        return new_tr

    for ci, chunk in enumerate(chunks):
        maxpipes = max((len(iv.get("Configurations") or []) for iv in chunk), default=1) or 1

        intervals_tr = clone_before(proto["intervals"])
        start_tr = clone_before(proto["start"])
        end_tr = clone_before(proto["end"])
        tubular_trs = [clone_before(proto_tubular) for _ in range(maxpipes)]
        interp_tr = clone_before(proto["interp"])
        response_tr = clone_before(proto["response"])
        offset_tr = clone_before(proto["offset"])

        base = ci * PER_BLOCK
        _fill_row(table, intervals_tr, LBL_INTERVALS,
                  [f"Interval {base + k + 1}" for k in range(len(chunk))], ncols)
        _fill_row(table, start_tr, LBL_START,
                  [_fmt_depth(iv.get("Start Depth (ft)")) for iv in chunk], ncols)
        _fill_row(table, end_tr, LBL_END,
                  [_fmt_depth(iv.get("End Depth (ft)")) for iv in chunk], ncols)
        for p, tr in enumerate(tubular_trs):
            vals = []
            for iv in chunk:
                cfgs = iv.get("Configurations") or []
                vals.append(cfgs[p] if p < len(cfgs) else "/")
            _fill_row(table, tr, LBL_TUBULAR, vals, ncols)
        _fill_row(table, interp_tr, LBL_INTERP,
                  ["-".join(str(c) for c in (iv.get("Channels") or [])) for iv in chunk], ncols)
        _fill_row(table, response_tr, LBL_RESPONSE,                        # from the Channels sheet, else blank
                  [str(iv.get("FirstResponse") or "") for iv in chunk], ncols)
        _fill_row(table, offset_tr, LBL_OFFSET,
                  ["/".join(str(o) for o in (iv.get("Offsets") or [])) for iv in chunk], ncols)

    for tr in proto["all"]:            # drop the prototype block
        tr.getparent().remove(tr)

    log(f"Interval table: {len(records)} interval(s) in {len(chunks)} block(s).")
    rev(f"Interval table — {len(records)} interval(s) written.")
    if own:
        doc.save(output_path)
    return len(records)

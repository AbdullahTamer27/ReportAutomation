"""Automation-report image placement.

Each image lives in a 1x1 single-cell table whose text is a tag (e.g. {{proc}}).
The tag is mapped to a filename, the file is inserted into that cell at a target
width, and if that would make it taller than MAX_HEIGHT it is capped by height
instead. Behavior is unchanged from the original script; it has only been turned
into a callable (`place_report_images`) with `print` routed through a callback.
"""

import os
import re
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from lxml import etree

# ---------------- Settings ----------------
# Target width for images (set to your image-cell width).
DEFAULT_IMG_WIDTH = Inches(6.46)
# No image may be taller than this. If filling width would exceed it,
# the image is capped to this height instead (width comes out proportional).
DEFAULT_MAX_HEIGHT = Inches(8.98)

# Static (config-level) image tags. Per-damage images are matched by pattern
# below, so any number of damage points scales without listing them here.
TAG_TO_FILE = {
    "{{proc}}":   "proc.jpg",
    "{{tempgr}}": "tempgr.jpg",
    "{{wh}}":     "wh.jpg",
    "{{raw}}":    "raw.jpg",
    "{{well}}":   "well.jpg",
    "{{ts}}":     "ts.jpg",
}

# Per-damage image tags: {{DMG<i>_<j>}} -> file "DMG<i>_<j>.<ext>".
DMG_TAG = re.compile(r"^\{\{DMG(\d+)_(\d+)\}\}$")


def _filename_for_tag(tag, tag_to_file):
    """Resolve a placeholder tag to a base filename (without forcing an
    extension). Static tags come from the dict; {{DMGi_j}} is pattern-matched."""
    if tag in tag_to_file:
        return tag_to_file[tag]
    m = DMG_TAG.match(tag)
    if m:
        return f"DMG{m.group(1)}_{m.group(2)}"
    return None


# Image extensions to accept, regardless of what's written in TAG_TO_FILE.
SUPPORTED_EXTS = (".png", ".jpg", ".jpeg", ".tiff", ".tif")


def _resolve_image_path(img_folder, fname):
    """Find the image for `fname` in `img_folder`, accepting any supported
    extension. Tries the configured name first, then the same base name with
    each supported extension (case-insensitive). Returns a path or None."""
    exact = os.path.join(img_folder, fname)
    if os.path.exists(exact):
        return exact
    stem = os.path.splitext(fname)[0]
    # Case-insensitive scan of the folder so e.g. PROC.JPG matches proc.jpg.
    try:
        entries = os.listdir(img_folder)
    except OSError:
        entries = []
    by_lower = {e.lower(): e for e in entries}
    for ext in SUPPORTED_EXTS:
        hit = by_lower.get((stem + ext).lower())
        if hit:
            return os.path.join(img_folder, hit)
    return None


# ---------------- Border helper ----------------
_NSMAP = {
    "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

def _add_image_border(pic_element, width_pt=1, color="000000"):
    """Append a solid border to an inline picture's shape properties."""
    spPr = pic_element.find(".//{%s}spPr" % _NSMAP["pic"])
    if spPr is None:
        return
    # Remove any existing border first
    for existing in spPr.findall("{%s}ln" % _NSMAP["a"]):
        spPr.remove(existing)
    width_emu = int(width_pt * 12700)
    ln = etree.SubElement(spPr, "{%s}ln" % _NSMAP["a"], w=str(width_emu))
    solidFill = etree.SubElement(ln, "{%s}solidFill" % _NSMAP["a"])
    etree.SubElement(solidFill, "{%s}srgbClr" % _NSMAP["a"], val=color)


# ---------------- Helper ----------------
def insert_image_gentle(cell, image_path, tag, img_width, max_height,
                        border_pt=0, border_color="000000"):
    """Insert the image without clearing the cell. Fill to img_width, but if that
    makes it taller than max_height, cap by height instead."""
    para = cell.paragraphs[0]
    for run in para.runs:
        if tag in run.text:
            run.text = run.text.replace(tag, "")
    if tag in "".join(r.text for r in para.runs):
        for run in para.runs:
            run.text = ""
    run = para.add_run()
    # First add at target width; python-docx scales height proportionally.
    pic = run.add_picture(image_path, width=img_width)
    # If the resulting height exceeds the cap, redo it sized by height.
    if pic.height > max_height:
        run.clear()
        pic = run.add_picture(image_path, height=max_height)
        sized_by = "height (capped)"
    else:
        sized_by = "width"
    if border_pt > 0:
        # pic._inline is the CT_Inline (<wp:inline>) element; the border helper
        # finds the descendant <pic:spPr> from it.
        _add_image_border(pic._inline, width_pt=border_pt, color=border_color)
    return sized_by


# ---------------- Alt-text placeholder pictures ----------------
# A template can mark an image slot with a real placeholder picture whose Alt
# Text (wp:docPr/@descr) holds the tag, e.g. {{proc}} or {{DMG1_1}}. We keep the
# picture's frame (size box, border, position) and only swap which image it
# points at, fitting the new image inside the box with its aspect preserved.
_TAG_TOKEN = re.compile(r"\{\{[^}]+\}\}")


def _alttext_tag(drawing):
    """Return the {{...}} tag found in a drawing's Alt Text (descr/name/title)."""
    docPr = drawing.find(".//" + qn("wp:docPr"))
    if docPr is None:
        return None
    for attr in ("descr", "name", "title"):
        v = docPr.get(attr)
        if v:
            m = _TAG_TOKEN.search(v)
            if m:
                return m.group(0)
    return None


def _fit_drawing_extent(drawing, image):
    """Resize the drawing's box to fit `image` inside it, preserving aspect.
    Updates both wp:extent (layout size) and the pic a:ext (shape size)."""
    extent = drawing.find(".//" + qn("wp:extent"))
    if extent is None or not image.px_height or not image.px_width:
        return
    try:
        box_cx, box_cy = int(extent.get("cx")), int(extent.get("cy"))
    except (TypeError, ValueError):
        return
    if box_cx <= 0 or box_cy <= 0:
        return
    aspect = image.px_width / image.px_height          # w/h
    if (box_cx / box_cy) > aspect:                     # box wider than image → fit height
        new_cy, new_cx = box_cy, int(round(box_cy * aspect))
    else:                                              # fit width
        new_cx, new_cy = box_cx, int(round(box_cx / aspect))
    extent.set("cx", str(new_cx)); extent.set("cy", str(new_cy))
    xfrm = drawing.find(".//" + qn("a:xfrm"))
    if xfrm is not None:
        a_ext = xfrm.find(qn("a:ext"))
        if a_ext is not None:
            a_ext.set("cx", str(new_cx)); a_ext.set("cy", str(new_cy))


def _replace_drawing_image(doc, drawing, image_path):
    """Re-point the drawing's picture at `image_path` (a new image part) and fit
    the box to the new image. Keeps the existing border / position."""
    blip = drawing.find(".//" + qn("a:blip"))
    if blip is None:
        return False
    rId, image = doc.part.get_or_add_image(image_path)
    blip.set(qn("r:embed"), rId)
    _fit_drawing_extent(drawing, image)
    return True


def place_images_by_alttext(doc, img_folder, tag_to_file, progress=None, review=None):
    """Fill every body picture whose Alt Text is a known image tag. Returns
    (placed, skipped, missing). Non-image tags (e.g. {{COMP}}) are ignored."""
    log = progress or print
    rev = review or (lambda m: None)
    placed, skipped, missing = 0, 0, []

    body = doc.element.body
    drawings = list(body.iter(qn("wp:inline"))) + list(body.iter(qn("wp:anchor")))
    for drawing in drawings:
        tag = _alttext_tag(drawing)
        if not tag:
            continue
        fname = _filename_for_tag(tag, tag_to_file)
        if not fname:
            continue   # tag present but not an image tag (e.g. {{COMP}}) — skip
        image_path = _resolve_image_path(img_folder, fname)
        if image_path:
            try:
                _replace_drawing_image(doc, drawing, image_path)
                log(f"OK placed {os.path.basename(image_path)} into {tag} (alt-text)")
                placed += 1
            except Exception as e:  # noqa: BLE001
                rev(f"❌ {tag}: image not placed — {os.path.basename(image_path)}: {e}")
                missing.append(os.path.basename(image_path)); skipped += 1
        else:
            stem = os.path.splitext(fname)[0]
            exts = "/".join(e.lstrip(".") for e in SUPPORTED_EXTS)
            rev(f"❌ {tag}: image not placed — no file '{stem}.[{exts}]' in folder")
            missing.append(fname); skipped += 1

    return placed, skipped, missing


# ---------------- Orchestration ----------------
def place_report_images(template_path, img_folder, output_path,
                        tag_to_file=None, img_width=DEFAULT_IMG_WIDTH,
                        max_height=DEFAULT_MAX_HEIGHT, progress=None,
                        border_pt=0.75, border_color="000000", review=None):
    """Place tagged images from `img_folder` into `template_path`, save to
    `output_path`. Returns {placed, skipped, missing}.

    `progress(msg)` streams verbose status; `review(msg)` streams only the
    curated review items (images not placed + reason, and a final summary)."""
    log = progress or print
    rev = review or (lambda m: None)
    if tag_to_file is None:
        tag_to_file = TAG_TO_FILE

    doc = Document(template_path)
    placed, skipped, missing = 0, 0, []

    for table in doc.tables:
        if len(table.rows) == 1 and len(table.columns) == 1:
            cell = table.rows[0].cells[0]
            tag = cell.text.strip()
            fname = _filename_for_tag(tag, tag_to_file)
            if fname:
                image_path = _resolve_image_path(img_folder, fname)
                if image_path:
                    try:
                        how = insert_image_gentle(cell, image_path, tag, img_width, max_height,
                                                 border_pt=border_pt, border_color=border_color)
                        log(f"OK placed {os.path.basename(image_path)} into {tag} (by {how})")
                        placed += 1
                    except Exception as e:  # noqa: BLE001
                        rev(f"❌ {tag}: image not placed — {os.path.basename(image_path)}: {e}")
                        missing.append(os.path.basename(image_path))
                        skipped += 1
                else:
                    stem = os.path.splitext(fname)[0]
                    exts = "/".join(e.lstrip(".") for e in SUPPORTED_EXTS)
                    rev(f"❌ {tag}: image not placed — no file '{stem}.[{exts}]' in folder")
                    missing.append(fname)
                    skipped += 1

    # Alt-text placeholder pictures (the table-free approach): swap in place.
    a_placed, a_skipped, a_missing = place_images_by_alttext(
        doc, img_folder, tag_to_file, progress=log, review=rev)
    placed += a_placed
    skipped += a_skipped
    missing += a_missing

    doc.save(output_path)
    rev(f"Images — placed {placed}, skipped {skipped}"
        + (f" (missing: {', '.join(missing)})" if missing else ""))
    log(f"Images: placed {placed}, skipped {skipped}. Saved -> {output_path}")
    return {"placed": placed, "skipped": skipped, "missing": missing}

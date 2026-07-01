"""Automation-report table filling.

Fills tables in a Word template from an Excel workbook, driven by tags placed
in each table's top-left header cell:

    {{joints_<sheet>}}   -> columns A..J (1..10)  of <sheet>
    {{highest_<sheet>}}  -> columns P..Y (16..25) of <sheet>, top N rows

Behavior is unchanged from the original script; it has only been turned into a
callable (`fill_report_tables`) so the report builder / UI can drive it, with
`print` routed through a `progress` callback.
"""

import re
import copy
import statistics

import openpyxl
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt

# ---------------- Settings ----------------
GRADE_COLORS = {"A": "FFFF00", "B": "FFC000", "C": "0070C0", "D": "FF0000"}
ROUND = {1: 1, 2: 1, 3: 1, 4: 3, 5: 3, 6: 1, 7: 1}

# Column layout (0-based) within each 10-col data block.
COLUMN_NAMES = [
    "#", "Top Body (ft)", "Bottom Body (ft)", "Body Length (ft)",
    "Nom Thk (in)", "Min Thk (in)", "Max Loss Depth (ft)",
    "Max Loss (%)", "Grade", "Damage Profile (% wall loss)",
]
JOINT_NO_IDX = 0
BODY_LEN_IDX = 3
MAX_LOSS_DEPTH_IDX = 6
MAX_LOSS_IDX = 7
GRADE_IDX = 8
DAMAGE_IDX = 9

# Damage Profile bar capping — stops long bars from widening/wrapping the cell.
DAMAGE_COL_WIDTH_IN = 2.05      # Word column width for the Damage Profile column
DAMAGE_CELL_MARGIN_IN = 0.08    # default Word side margin, per side
BAR_FONT_PT = 10.0              # font size of the bar text
BAR_CHAR_EM = 0.6              # bar glyph width as a fraction of the font size
                               # (0.6 ≈ a full block █, the widest common case)


def _max_bar_chars():
    """Max bar characters that fit the Damage Profile column width."""
    usable_in = DAMAGE_COL_WIDTH_IN - 2 * DAMAGE_CELL_MARGIN_IN
    char_in = (BAR_FONT_PT / 72.0) * BAR_CHAR_EM
    return max(1, int(usable_in // char_in))


MAX_BAR_CHARS = _max_bar_chars()

# Joint (body) length is checked RELATIVE to the string's own typical (median)
# length, not a fixed range — so a 28 ft tubing or a 60 ft casing string doesn't
# fire wholesale; only joints that deviate from their neighbours are flagged.
# The band is asymmetric: joints legitimately run short (pup / shoe joints), so
# allow more slack below the median; an over-length joint is more suspicious.
BODY_LEN_SHORT_TOL = 0.15      # flag if > 15% SHORTER than the string median
BODY_LEN_LONG_TOL = 0.10       # flag if > 10% LONGER than the string median
BODY_LEN_MIN_SAMPLE = 4        # need at least this many joints to trust a median


def grade_for_loss(loss):
    """Grade from Max Loss (%): 0–4.9 A, 5–9.9 B, 10–19.9 C, 20+ D."""
    if loss < 5:
        return "A"
    if loss < 10:
        return "B"
    if loss < 20:
        return "C"
    return "D"


def _joint_label(vals):
    jn = vals[JOINT_NO_IDX]
    if isinstance(jn, float) and jn.is_integer():
        jn = int(jn)
    return jn


def typical_body_length(data_rows):
    """The string's typical joint length = median of its real joints' body
    lengths. Returns None if there are too few to establish a norm."""
    lengths = [
        v[BODY_LEN_IDX] for v in data_rows
        if not is_excluded(v)
        and isinstance(v[BODY_LEN_IDX], (int, float)) and v[BODY_LEN_IDX] > 0
    ]
    if len(lengths) < BODY_LEN_MIN_SAMPLE:
        return None
    return statistics.median(lengths)


def review_row(table_name, vals, review, typical_len=None):
    """Sanity-check one real (non-excluded) data row. Corrects vals[GRADE_IDX]
    in place when it disagrees with Max Loss (%). Emits messages via review(msg).

    `typical_len` (the string's median joint length) enables the body-length
    check; when omitted (e.g. the highest-loss table, which the joints table
    already covers) the length check is skipped."""
    if review is None:
        return
    jn = _joint_label(vals)

    # Any negative number in the row.
    for i, v in enumerate(vals):
        if isinstance(v, (int, float)) and v < 0:
            review(f"⚠ {table_name} joint {jn}: negative {COLUMN_NAMES[i]} ({v})")

    # Body length vs the string's own typical length — asymmetric band.
    bl = vals[BODY_LEN_IDX]
    if typical_len and isinstance(bl, (int, float)) and bl > 0:
        low = typical_len * (1 - BODY_LEN_SHORT_TOL)
        high = typical_len * (1 + BODY_LEN_LONG_TOL)
        if bl < low or bl > high:
            review(f"⚠ {table_name} joint {jn}: Body Length {bl:.1f} ft is off this "
                   f"string's typical {typical_len:.1f} ft (expected {low:.1f}–{high:.1f})")

    # Grade vs Max Loss (%) — correct the grade, never the value.
    loss = vals[MAX_LOSS_IDX]
    if isinstance(loss, (int, float)) and loss >= 0:
        if loss > 100:
            review(f"⚠ {table_name} joint {jn}: Max Loss {loss:.1f}% exceeds 100%")
        correct = grade_for_loss(loss)
        current = str(vals[GRADE_IDX]).strip().upper()
        if current != correct:
            review(f"✎ {table_name} joint {jn}: grade {current or '—'}→{correct} "
                   f"(Max Loss {loss:.1f}%) — corrected")
        # Always write the canonical (uppercase) grade derived from the loss, even
        # when it only differed in case/spacing — otherwise the cell keeps its odd
        # form and its colour (looked up on uppercase A–D) is never applied.
        vals[GRADE_IDX] = correct

JOINTS_TAG = re.compile(r"\{\{joints_(\w+)\}\}")
HIGHEST_TAG = re.compile(r"\{\{highest_(\w+)\}\}")
SUMMARY_TAG = "{{SUMMARY}}"

HIGHEST_TOP_N = 4   # how many rows in the highest-loss tables


# ---------------- Excel reading ----------------
def read_table_block(ws, cols, start_col_for_hashcheck):
    """Read rows for a set of 1-based column indices. cols is a list of 10 column
    indices (the A..J-equivalent block). Stops when the '#' column stops being numeric."""
    rows = []
    hash_col = cols[0]
    for r in range(2, ws.max_row + 1):
        a = ws.cell(row=r, column=hash_col).value
        if not isinstance(a, (int, float)):
            if rows:
                break
            continue
        rows.append([ws.cell(row=r, column=c).value for c in cols])
    return rows


def read_joints(ws):
    # columns A..J = 1..10
    return read_table_block(ws, list(range(1, 11)), 1)


def read_highest(ws, top_n):
    # columns P..Y = 16..25  (P=#, Q..Y = the 10-col block; Z=rank ignored)
    # The block is pre-ranked by severity, so the worst joints sit at the top.
    rows = read_table_block(ws, list(range(16, 26)), 16)

    # Drop rows where Max Loss (%) is 0 or negative — not real measurements.
    # Annotated/excluded rows (non-numeric Max Loss) are kept as-is.
    valid_rows = [
        v for v in rows
        if not isinstance(v[MAX_LOSS_IDX], (int, float)) or v[MAX_LOSS_IDX] > 0
    ]

    # If fewer valid rows than top_n, show only what's available — no padding.
    # If more than top_n exist and some are C/D, extend to include all C/D rows.
    cd_count = sum(
        1 for v in valid_rows
        if isinstance(v[MAX_LOSS_IDX], (int, float)) and str(v[GRADE_IDX]).strip() in ("C", "D")
    )
    n = max(min(top_n, len(valid_rows)), cd_count)
    return valid_rows[:n]


def is_excluded(vals):
    """An annotated joint: the Max Loss (%) column holds a text note instead of a
    number (or reads blank because the cells were merged in Excel). Either way the
    measurement columns are merged into one note cell in the Word table."""
    return not isinstance(vals[MAX_LOSS_IDX], (int, float))


def annotation_text(vals):
    """The note shown in an annotated joint's merged cell. Prefer the Max Loss (%)
    column (where it lives when Excel isn't merged); fall back to the first text
    found across the merged span (where it lives when Excel is already merged)."""
    note = vals[MAX_LOSS_IDX]
    if isinstance(note, str) and note.strip():
        return note.strip()
    for i in range(5, 10):
        v = vals[i]
        if isinstance(v, str) and v.strip():
            return v.strip()
    return ""


def fmt(val, idx):
    if val is None:
        return ""
    if isinstance(val, (int, float)) and idx in ROUND:
        return f"{val:.{ROUND[idx]}f}"
    return str(val)


# ---------------- Word helpers ----------------
def hex_to_rgb(h):
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def clear_cell_shading(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)


def set_cell_bg(cell, hexc):
    clear_cell_shading(cell)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexc)
    tcPr.append(shd)


def set_cell_font_color(cell, hexc):
    rgb = hex_to_rgb(hexc)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = rgb


def set_cell_text(cell, text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(text)


def clear_cell_paragraphs(cell):
    for p in cell.paragraphs[1:]:
        p._p.getparent().remove(p._p)
    first = cell.paragraphs[0]
    for run in first.runs:
        run.text = ""
    return first


def reset_cell(cell):
    set_cell_text(cell, "")
    clear_cell_shading(cell)


def clone_row(table, template_row):
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def fill_table(table, data_rows, table_name=None, review=None):
    """Fill a tagged table with data_rows using the cloned-style-row approach.
    Real data rows are sanity-checked (and grade-corrected) via review_row."""
    template_row = table.rows[1]   # the single styled data row
    typical_len = typical_body_length(data_rows)   # string's median joint length
    for vals in data_rows:
        new_row = clone_row(table, template_row)
        cells = new_row.cells
        for c in cells:
            reset_cell(c)
        if is_excluded(vals):
            # Annotated joint: keep the geometry columns (0–4), merge the
            # measurement columns (5–9) into one cell and show the note. This
            # merge is done here in Word whether or not the Excel was merged.
            for i in range(5):
                set_cell_text(cells[i], fmt(vals[i], i))
            merged = cells[5]
            for i in range(6, 10):
                merged = merged.merge(cells[i])
            para = clear_cell_paragraphs(merged)
            run = para.add_run(annotation_text(vals))
            run.font.name = "Calibri"; run.font.size = Pt(10)
        else:
            # Review + correct grade before writing the cells.
            review_row(table_name, vals, review, typical_len=typical_len)
            for i in range(10):
                text = fmt(vals[i], i)
                # Cap the Damage Profile bar so a long bar can't widen/wrap the cell.
                if i == DAMAGE_IDX and isinstance(vals[i], str) and len(text) > MAX_BAR_CHARS:
                    text = text[:MAX_BAR_CHARS]
                set_cell_text(cells[i], text)
            grade = str(vals[8]).strip()
            if grade in GRADE_COLORS:
                set_cell_bg(cells[8], GRADE_COLORS[grade])
                set_cell_font_color(cells[9], GRADE_COLORS[grade])
    # remove the template row
    template_row._tr.getparent().remove(template_row._tr)


def delete_table(table):
    table._tbl.getparent().remove(table._tbl)
    # NOTE: this leaves the heading/caption around the table behind.
    # Pending the answer on how headings are structured, we'll also remove those.


# ---------------- Summary table (worst joint per pipe) ----------------
def _strip_tag(cell, tag):
    """Remove `tag` from a cell, touching only the runs it spans (so the rest of
    the header label keeps its text/formatting). Handles a tag split across runs."""
    for para in cell.paragraphs:
        ts = para.runs
        joined = "".join(r.text for r in ts)
        if tag not in joined:
            continue
        # Fast path: a single run holds the whole tag.
        if any(tag in r.text for r in ts):
            for r in ts:
                if tag in r.text:
                    r.text = r.text.replace(tag, "")
            continue
        # Split across runs: drop the tag chars from the spanning runs only.
        idx = joined.find(tag)
        end = idx + len(tag)
        pos = 0
        for r in ts:
            seg_start, seg_end = pos, pos + len(r.text)
            pos = seg_end
            if seg_end <= idx or seg_start >= end:
                continue
            left = r.text[: max(0, idx - seg_start)]
            right = r.text[end - seg_start:] if end <= seg_end else ""
            r.text = left + right


def worst_joint(ws, top_n, table_name=None, review=None):
    """The worst (highest metal loss) real joint for a pipe sheet, or None.

    Uses the same severity-ranked rows as the highest-loss table and returns the
    first row that is a genuine measurement (numeric Max Loss, not an annotated/
    excluded joint). Grade is corrected against Max Loss via review_row, exactly
    like the per-pipe tables."""
    for vals in read_highest(ws, top_n):
        if is_excluded(vals) or not isinstance(vals[MAX_LOSS_IDX], (int, float)):
            continue
        review_row(table_name, vals, review)   # corrects vals[GRADE_IDX] in place
        return vals
    return None


def _fill_summary_row(row, vals, rev, suffix=None):
    """Write a summary data row: optional col-0 pipe suffix, then Max Loss (%),
    Grade (+ colour), Max Loss Depth (ft) from the pipe's worst joint `vals`
    (which may be None — then only the suffix is written)."""
    cells = row.cells
    if suffix is not None:
        set_cell_text(cells[0], suffix)                                    # pipe name
    if vals is None:
        return
    set_cell_text(cells[1], fmt(vals[MAX_LOSS_IDX], MAX_LOSS_IDX))         # metal loss
    grade = str(vals[GRADE_IDX]).strip()
    set_cell_text(cells[2], grade)                                         # grade
    if grade in GRADE_COLORS:
        set_cell_bg(cells[2], GRADE_COLORS[grade])                         # + bg colour
    set_cell_text(cells[3], fmt(vals[MAX_LOSS_DEPTH_IDX], MAX_LOSS_DEPTH_IDX))


def fill_summary_table(table, wb, sheets, pipe_order, highest_top_n,
                       progress=None, review=None, pipe_model=None):
    """Fill the cross-pipe summary table in place.

    Columns: 0 = pipe name, 1 = Max Loss (%), 2 = Grade (+ colour), 3 = Max Loss
    Depth (ft). Row mapping is bottom-anchored: the FIRST pipe fills the LAST data
    row, the second the second-to-last, and so on.

    Universal mode (`pipe_model` given): column 0 is filled with each pipe's
    `suffix`, and the unused top rows are deleted (master template has a fixed 7
    rows). Legacy mode: column 0 is left untouched and no rows are removed."""
    log = progress or print
    rev = review or (lambda m: None)

    data_rows = table.rows[1:]   # row 0 is the header (holds the {{SUMMARY}} tag)
    nrows = len(data_rows)
    filled = 0

    if pipe_model is not None:
        # Dynamic, like the other tables: the template has 1 header + 1 styled
        # data row, which we clone once per pipe. Cloning appends to the bottom,
        # so iterating in reverse puts the FIRST pipe in the LAST row.
        if len(table.rows) < 2:
            rev("⚠ Summary: needs a header row + one data row to clone — nothing filled.")
            _strip_tag(table.rows[0].cells[0], SUMMARY_TAG)
            return 0
        template_row = table.rows[1]
        for p in reversed(pipe_model):
            new_row = clone_row(table, template_row)
            for c in new_row.cells:
                reset_cell(c)
            sheet_name = p["sheet"]
            vals = None
            if sheet_name in sheets:
                vals = worst_joint(wb[sheet_name], highest_top_n,
                                   table_name=f"summary[{sheet_name}]", review=rev)
                if vals is None:
                    rev(f"⚠ Summary: '{sheet_name}' has no valid joint — data left blank")
            else:
                rev(f"⚠ Summary: sheet '{sheet_name}' not in workbook — data left blank")
            _fill_summary_row(new_row, vals, rev, suffix=p.get("suffix", ""))
            filled += 1
        template_row._tr.getparent().remove(template_row._tr)   # drop the template row
    else:
        if nrows != len(pipe_order):
            rev(f"⚠ Summary: table has {nrows} data row(s) but {len(pipe_order)} "
                f"pipe(s) — filling the last {min(nrows, len(pipe_order))}.")
        for i, sheet_name in enumerate(pipe_order):
            ri = nrows - 1 - i        # first pipe → last row
            if ri < 0:
                break
            if sheet_name not in sheets:
                rev(f"⚠ Summary: sheet '{sheet_name}' not in workbook — row left blank")
                continue
            vals = worst_joint(wb[sheet_name], highest_top_n,
                               table_name=f"summary[{sheet_name}]", review=rev)
            if vals is None:
                rev(f"⚠ Summary: '{sheet_name}' has no valid joint — row left blank")
                continue
            _fill_summary_row(data_rows[ri], vals, rev, suffix=None)
            filled += 1

    _strip_tag(table.rows[0].cells[0], SUMMARY_TAG)
    log(f"OK summary: {filled} pipe(s) filled")
    return filled


# ---------------- Orchestration ----------------
def fill_report_tables(template_path, workbook_path, output_path,
                       highest_top_n=HIGHEST_TOP_N, progress=None, review=None,
                       pipe_model=None):
    """Fill all tagged tables in `template_path` from `workbook_path` and save to
    `output_path`. Returns a result dict: {filled, deleted, used, warnings}.

    `progress(msg)` streams verbose status; `review(msg)` streams only the
    curated review items (failures, warnings, data-sanity flags)."""
    log = progress or print
    rev = review or (lambda m: None)

    wb = openpyxl.load_workbook(workbook_path, data_only=True)
    sheets = set(wb.sheetnames)

    doc = Document(template_path)

    # Snapshot the tables up front (we'll be modifying the doc)
    tables = list(doc.tables)

    filled, deleted, warnings = 0, 0, []
    used = set()
    # Pipes in document order (shallow→deep): the order their joints/highest tags
    # first appear. The summary table is filled positionally against this list.
    pipe_order = []
    summary_tables = []

    def _note_pipe(sheet_name):
        if sheet_name not in pipe_order:
            pipe_order.append(sheet_name)

    for table in tables:
        if not table.rows:
            continue
        header0 = table.rows[0].cells[0].text

        # Summary table is only 4 columns, so handle (stash) it before the
        # >=10-column guard. Filled after the loop, once pipe_order is complete.
        if SUMMARY_TAG in header0:
            summary_tables.append(table)
            continue

        if len(table.rows) < 2 or len(table.columns) < 10:
            continue

        mj = JOINTS_TAG.search(header0)
        mh = HIGHEST_TAG.search(header0)

        if mj:
            sheet_name = mj.group(1)
            tag = f"joints_{sheet_name}"
            if sheet_name in sheets:
                try:
                    rows = read_joints(wb[sheet_name])
                    fill_table(table, rows, table_name=tag, review=rev)
                    set_cell_text(table.rows[0].cells[0], "#")  # restore header
                    used.add(sheet_name)
                    _note_pipe(sheet_name)
                    filled += 1
                    log(f"OK {tag}: {len(rows)} rows")
                except Exception as e:  # noqa: BLE001
                    rev(f"❌ {tag}: failed to fill table — {e}")
            else:
                delete_table(table)
                deleted += 1
                rev(f"⚠ {tag}: sheet not found in workbook → table removed")

        elif mh:
            sheet_name = mh.group(1)
            tag = f"highest_{sheet_name}"
            if sheet_name in sheets:
                try:
                    rows = read_highest(wb[sheet_name], highest_top_n)
                    fill_table(table, rows, table_name=tag, review=rev)
                    set_cell_text(table.rows[0].cells[0], "#")
                    _note_pipe(sheet_name)
                    filled += 1
                    log(f"OK {tag}: {len(rows)} rows")
                except Exception as e:  # noqa: BLE001
                    rev(f"❌ {tag}: failed to fill table — {e}")
            else:
                delete_table(table)
                deleted += 1
                rev(f"⚠ {tag}: sheet not found in workbook → table removed")

    # Summary table(s): worst joint per pipe. With a pipe_model (universal master
    # template) column 0 = pipe suffix and unused rows are trimmed; otherwise the
    # legacy bottom-anchored fill keyed by tag-appearance order.
    summary_order = [p["sheet"] for p in pipe_model] if pipe_model else pipe_order
    for st in summary_tables:
        try:
            fill_summary_table(st, wb, sheets, summary_order, highest_top_n,
                               progress=log, review=rev, pipe_model=pipe_model)
            filled += 1
        except Exception as e:  # noqa: BLE001
            rev(f"❌ summary: failed to fill table — {e}")

    # Validation: pipe sheets that exist but were never filled by any tag
    pipe_sheets = {s for s in sheets if s.endswith("Pipe")}
    for s in pipe_sheets - used:
        warnings.append(f"sheet '{s}' has no matching tag in template")
        rev(f"⚠ Sheet '{s}' exists but has no matching tag in the template")

    doc.save(output_path)
    log(f"Tables: filled {filled}, deleted {deleted}. Saved -> {output_path}")

    return {"filled": filled, "deleted": deleted, "used": used, "warnings": warnings}

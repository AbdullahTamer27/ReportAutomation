"""Unit tests for damage-block expansion — the 0 / 1 / N logic, directly.

These don't generate full reports; they drive ``expand_damage_blocks`` on a tiny
in-memory document, so the "choose 0 damages vs 1 vs many" behaviour is pinned
precisely and cheaply, independent of any golden fixture.
"""

import re

import pytest
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

from well_tools.report.damage_blocks import expand_damage_blocks

_NS = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
       'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
       'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
       'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
       'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"')


def _doc_with_block(include_textbox=False):
    """A document with one damage block: an image tag and (optionally) an
    overlay text box, both using the @N sentinel."""
    doc = Document()
    body = doc.element.body
    sect = body.find(qn("w:sectPr"))

    def addp(text):
        sect.addprevious(parse_xml(
            f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:r><w:t>{text}</w:t></w:r></w:p>'))

    addp("{{damage_block_start}}")
    addp("Photo {{DMG@N_1}}")
    if include_textbox:
        sect.addprevious(parse_xml(
            f'<w:p {_NS}><w:r><mc:AlternateContent><mc:Choice Requires="wps"><w:drawing>'
            f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="1" '
            f'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
            f'<wp:simplePos x="0" y="0"/><wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>'
            f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>'
            f'<wp:extent cx="900000" cy="300000"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
            f'<wp:docPr id="1" name="tb"/><a:graphic>'
            f'<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
            f'<wps:wsp><wps:cNvSpPr txBox="1"/><wps:spPr/><wps:txbx><w:txbxContent>'
            f'<w:p><w:r><w:t>{{{{ovl_ml@N_1}}}}</w:t></w:r></w:p></w:txbxContent></wps:txbx>'
            f'<wps:bodyPr/></wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing>'
            f'</mc:Choice></mc:AlternateContent></w:r></w:p>'))
    addp("{{damage_block_end}}")
    return doc


def _body_text(doc):
    return "".join(t.text or "" for t in doc.element.body.iter(qn("w:t")))


def _dmg_ids(doc):
    return sorted(set(re.findall(r"\{\{DMG(\d+)_1\}\}", _body_text(doc))))


@pytest.mark.parametrize("count,expected_ids", [
    (0, []),
    (1, ["1"]),
    (3, ["1", "2", "3"]),
])
def test_expand_count(count, expected_ids):
    doc = _doc_with_block()
    found = expand_damage_blocks(doc, count)
    assert found is True                      # markers were located
    assert _dmg_ids(doc) == expected_ids
    # markers themselves are always consumed
    assert "damage_block_start" not in _body_text(doc)
    assert "damage_block_end" not in _body_text(doc)


def test_zero_removes_block_entirely():
    doc = _doc_with_block()
    expand_damage_blocks(doc, 0)
    assert "DMG" not in _body_text(doc)


def test_missing_markers_returns_false():
    doc = Document()
    doc.add_paragraph("no markers here")
    assert expand_damage_blocks(doc, 2) is False


def test_at_n_substituted_inside_textboxes():
    """Regression: @N must renumber inside the overlay text box (both the
    DrawingML and VML copies), not double up or blank one out."""
    doc = _doc_with_block(include_textbox=True)
    expand_damage_blocks(doc, 2)
    box_texts = ["".join(t.text or "" for t in txbx.iter(qn("w:t")))
                 for txbx in doc.element.body.iter(qn("w:txbxContent"))]
    # Two blocks → the tag renumbers to _1 and _2; never left as @N, never doubled.
    assert "{{ovl_ml1_1}}" in box_texts
    assert "{{ovl_ml2_1}}" in box_texts
    assert all("@N" not in t for t in box_texts)
    assert all(t.count("ovl_ml") <= 1 for t in box_texts)   # not doubled

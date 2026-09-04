"""Template introspection (Epic C3): tag scan + per-template field selection."""

from docx import Document

from webapp.introspect import extract_tags, template_form


def _make(tmp_path, body_runs=None, cell_text=None, header_text=None):
    doc = Document()
    if body_runs:
        p = doc.add_paragraph()
        for r in body_runs:           # multiple runs → tests split-across-runs joining
            p.add_run(r)
    if cell_text is not None:
        doc.add_table(rows=1, cols=1).rows[0].cells[0].text = cell_text
    if header_text is not None:
        doc.sections[0].header.paragraphs[0].text = header_text
    path = str(tmp_path / "t.docx")
    doc.save(path)
    return path


def test_extract_tags_body_table_header_and_split_runs(tmp_path):
    p = _make(
        tmp_path,
        body_runs=["Well {{well_name}} and ", "{{orig", "_comp}}"],   # split tag
        cell_text="Loss {{SUMMARY}}",
        header_text="{{COMPNAME}} report",
    )
    tags = extract_tags(p)
    assert "{{well_name}}" in tags
    assert "{{orig_comp}}" in tags        # reassembled across runs
    assert "{{SUMMARY}}" in tags          # from the table cell
    assert "{{COMPNAME}}" in tags         # from the header


def test_template_fields_filters_and_orders(tmp_path):
    # A template using: two user tags, engine/derived tags, and one unknown tag.
    p = _make(tmp_path, body_runs=[
        "{{log_date}} {{well_name}} {{SUMMARY}} {{pie_firstPipe}} "
        "{{firstPipe_name}} {{casings}} {{tool_type}} {{block}}"
    ])
    fields = template_form(p)["fields"]
    keys = [f["key"] for f in fields]
    # registry user fields present, in registry order (well_name before log_date),
    # then the unknown user tag as a generic text box.
    assert keys == ["well_name", "log_date", "block"]
    # engine/derived tags never become fields
    assert not any(k in keys for k in
                   ("SUMMARY", "pie_firstPipe", "firstPipe_name", "casings", "tool_type"))
    # the generic field is a plain text input labelled from the tag
    block = next(f for f in fields if f["key"] == "block")
    assert block["tag"] == "{{block}}" and block["type"] == "text" and block["label"] == "Block"


def test_template_with_no_metadata_tags_yields_no_fields(tmp_path):
    p = _make(tmp_path, body_runs=["{{SUMMARY}} {{pie_thirdPipe}} {{DMG1_1}} {{INTERVALS}}"])
    assert template_form(p)["fields"] == []


def test_saudi_style_template_shows_all_seven(tmp_path):
    p = _make(tmp_path, body_runs=[
        "{{well_name}} {{field}} {{well_type}} {{btm_depth}} "
        "{{log_date}} {{orig_comp}} {{last_wko}}"
    ])
    keys = [f["key"] for f in template_form(p)["fields"]]
    assert keys == ["well_name", "field", "well_type", "btm_depth",
                    "log_date", "orig_comp", "last_wko"]


def _controls(form):
    return {c["key"]: c["present"] for c in form["controls"]}


def test_controls_gated_on_their_tags(tmp_path):
    # Template drives only the disclaimer and FW16 (tool_type); no damage / wellhead.
    p = _make(tmp_path, body_runs=["{{well_name}} {{DISC}} {{tool_type}}"])
    ctl = _controls(template_form(p))
    assert ctl == {"damage_count": False, "include_disclaimer": True,
                   "wellhead_damage": False, "fw16": True}


def test_all_controls_when_template_uses_them(tmp_path):
    p = _make(tmp_path, body_runs=[
        "{{DISC}} {{ovl_wellhead}} {{tool_type}} {{damage_block_start}}"])
    assert all(template_form(p)["controls"][i]["present"] for i in range(4))


def test_fast_scan_matches_python_docx_reference(tmp_path):
    # The zip+regex scanner must agree with the python-docx walk on the tricky
    # cases: tags split across runs, in a table cell, and in a header.
    from webapp.introspect import _scan_zip, _scan_docx
    p = _make(
        tmp_path,
        body_runs=["{{well_name}} ", "{{orig", "_comp}} ", "{{pie_firstPipe}}"],
        cell_text="{{SUMMARY}} {{log_date}}",
        header_text="{{COMPNAME}}",
    )
    assert _scan_zip(p) == _scan_docx(p)


def test_extract_tags_caches_and_invalidates_on_change(tmp_path):
    from webapp.introspect import extract_tags
    p = _make(tmp_path, body_runs=["{{well_name}}"])
    assert extract_tags(p) == {"{{well_name}}"}
    assert extract_tags(p) == {"{{well_name}}"}          # served from cache

    # Rewrite the file with different tags → new mtime/size → rescanned.
    import time
    time.sleep(0.01)
    doc = Document()
    doc.add_paragraph("{{field}} {{block}}")
    doc.save(p)
    assert extract_tags(p) == {"{{field}}", "{{block}}"}


def test_cached_result_is_not_mutable_by_callers(tmp_path):
    from webapp.introspect import extract_tags
    p = _make(tmp_path, body_runs=["{{well_name}}"])
    got = extract_tags(p)
    got.add("{{injected}}")                               # mutating the copy…
    assert "{{injected}}" not in extract_tags(p)          # …must not poison the cache


def test_alt_text_tags_are_found(tmp_path):
    """A picture's tag lives in its Alt Text, which Word stores as an attribute
    of the drawing — not as text. Stripping XML tags takes their attributes with
    them, so those tags were invisible to the scan, and anything gated on one
    ({{ops}}) would never fire."""
    from docx import Document
    from docx.shared import Inches
    from PIL import Image

    Image.new("RGB", (40, 20), (0, 0, 0)).save(str(tmp_path / "x.png"))
    doc = Document()
    doc.add_paragraph("{{well_name}}")
    run = doc.add_paragraph().add_run()
    run.add_picture(str(tmp_path / "x.png"), width=Inches(1))
    doc.inline_shapes[0]._inline.docPr.set("descr", "{{ops}}")
    path = str(tmp_path / "t.docx")
    doc.save(path)

    tags = extract_tags(path)
    assert "{{ops}}" in tags
    assert "{{well_name}}" in tags


def test_the_scan_does_not_match_across_split_runs(tmp_path):
    """Word splits a tag's text across runs, so the body is scanned with the
    markup removed. Scanning the raw XML instead would let the pattern run from
    one run's "{{" to another's "}}" and match a mouthful of XML as a tag."""
    from docx import Document

    doc = Document()
    para = doc.add_paragraph()
    para.add_run("{{well")          # deliberately split, as Word does
    para.add_run("_name}}")
    para.add_run(" and {{field}}")
    path = str(tmp_path / "split.docx")
    doc.save(path)

    tags = extract_tags(path)
    assert tags == {"{{well_name}}", "{{field}}"}
    assert not any("<" in t or len(t) > 40 for t in tags)

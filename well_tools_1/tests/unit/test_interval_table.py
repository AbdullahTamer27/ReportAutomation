"""Unit tests for the dynamic interval-table pass."""

from docx import Document

from well_tools.report import interval_table as it


def _prototype_doc():
    """A minimal template: two meta rows + a single {{INTERVALS}} prototype block."""
    doc = Document()
    labels = [
        ("Well Name", ""), ("RIG Name", "RIGLESS"),
        (it.INTERVALS_TAG, ""),          # the Intervals prototype row (tagged)
        ("Start Depth(ft)", ""), ("End Depth(ft)", ""),
        ("Tubular size & weight", ""),   # one tubular prototype row
        ("Interpretation Channels", ""), ("Pipe channel response", ""), ("Offset", ""),
    ]
    table = doc.add_table(rows=len(labels), cols=4)
    for r, (c0, c1) in enumerate(labels):
        table.rows[r].cells[0].text = c0
        table.rows[r].cells[1].text = c1
    return doc, table


def _grid(table):
    return [[c.text for c in r.cells] for r in table.rows]


def _find(grid, label):
    return [row for row in grid if row[0] == label]


def _records():
    # 5 intervals -> chunks of 3: block1 = [4,3,2] pipes, block2 = [3,1] pipes.
    return [
        {"Start Depth (ft)": 5, "End Depth (ft)": 117, "Configurations": ["A", "B", "C", "D"],
         "Channels": [1, 2, 3, 4], "Offsets": [0.1, 0.2, 0.3, 0.4]},
        {"Start Depth (ft)": 117, "End Depth (ft)": 307, "Configurations": ["A", "B", "C"],
         "Channels": [1, 2, 3], "Offsets": [0.1, 0.2, 0.3]},
        {"Start Depth (ft)": 307, "End Depth (ft)": 510.5, "Configurations": ["A", "B"],
         "Channels": [1, 2], "Offsets": [0.1, 0.2]},
        {"Start Depth (ft)": 510.5, "End Depth (ft)": 2234, "Configurations": ["A", "B", "C"],
         "Channels": [1, 2, 3], "Offsets": [0.1, 0.2, 0.3]},
        {"Start Depth (ft)": 2234, "End Depth (ft)": 7233, "Configurations": ["A"],
         "Channels": [1], "Offsets": [0.1]},
    ]


def test_fill_full():
    doc, table = _prototype_doc()
    n = it.place_interval_table(None, _records(), well_name="ZULF_65", doc=doc)
    assert n == 5
    grid = _grid(table)

    # tag gone; meta rows correct (RIG Name untouched, Well Name filled)
    assert not any(it.INTERVALS_TAG in c for row in grid for c in row)
    assert _find(grid, "Well Name")[0][1] == "ZULF_65"
    assert _find(grid, "RIG Name")[0][1] == "RIGLESS"

    # two blocks
    ivs = _find(grid, "Intervals")
    assert len(ivs) == 2
    assert ivs[0][1:] == ["Interval 1", "Interval 2", "Interval 3"]
    assert ivs[1][1:] == ["Interval 4", "Interval 5", ""]     # last chunk has 2 intervals

    # depths (whole numbers bare, decimals kept)
    assert _find(grid, "Start Depth(ft)")[0][1:] == ["5", "117", "307"]
    assert _find(grid, "End Depth(ft)")[0][1:] == ["117", "307", "510.5"]

    # block 1 = 4 tubular rows; the 4th pipe row has "/" where an interval has < 4 pipes
    tub = _find(grid, "Tubular size & weight")
    assert len(tub) == 4 + 3                                   # block1 4 + block2 3
    assert tub[3][1:] == ["D", "/", "/"]                       # pipe 4: only interval 1 has it

    # interpretation channels joined with "-"; response blank; offset joined with "/"
    assert _find(grid, "Interpretation Channels")[0][1:] == ["1-2-3-4", "1-2-3", "1-2"]
    assert _find(grid, "Pipe channel response")[0][1:] == ["", "", ""]
    assert _find(grid, "Offset")[0][1] == "0.1/0.2/0.3/0.4"


def test_no_thickness_leaves_channel_rows_blank():
    doc, table = _prototype_doc()
    recs = [{"Start Depth (ft)": 5, "End Depth (ft)": 117, "Configurations": ["A", "B"]}]
    it.place_interval_table(None, recs, well_name="W", doc=doc)
    grid = _grid(table)
    assert _find(grid, "Interpretation Channels")[0][1:] == ["", "", ""]
    assert _find(grid, "Pipe channel response")[0][1:] == ["", "", ""]
    assert _find(grid, "Offset")[0][1:] == ["", "", ""]
    # the table still builds with the pipe rows present
    assert _find(grid, "Tubular size & weight")[0][1:] == ["A", "", ""]


def test_no_interval_table_is_a_noop():
    doc = Document()
    doc.add_paragraph("no interval table here")
    assert it.place_interval_table(None, _records(), doc=doc) == 0


def test_read_first_response(tmp_path):
    from openpyxl import Workbook
    from openpyxl.utils import column_index_from_string
    wb = Workbook()
    ws = wb.active
    ws.title = it.FIRST_RESPONSE_SHEET                 # "Channels"
    col = column_index_from_string(it.FIRST_RESPONSE_COL)
    for i, v in enumerate(["10-20-30-40", "12-33-44-70", 1235]):
        ws.cell(row=it.FIRST_RESPONSE_START_ROW + i, column=col, value=v)
    p = str(tmp_path / "data.xlsx")
    wb.save(p)
    assert it._read_first_response(p, 3) == ["10-20-30-40", "12-33-44-70", "1235"]
    assert it._read_first_response(p, 5)[3:] == ["", ""]   # more asked than present


def test_read_first_response_absent(tmp_path):
    from openpyxl import Workbook
    wb = Workbook(); wb.active.title = "Other"
    p = str(tmp_path / "d.xlsx"); wb.save(p)
    assert it._read_first_response(p, 3) == ["", "", ""]   # no Channels sheet
    assert it._read_first_response(None, 2) == ["", ""]     # no workbook


def test_pipe_channel_response_filled_from_records():
    doc, table = _prototype_doc()
    recs = _records()
    for r, v in zip(recs, ["10-20-30-40", "12-33-44", "12-35", "9-9-9", "7"]):
        r["FirstResponse"] = v
    it.place_interval_table(None, recs, well_name="W", doc=doc)
    grid = _grid(table)
    resp = _find(grid, "Pipe channel response")
    assert resp[0][1:] == ["10-20-30-40", "12-33-44", "12-35"]   # block 1
    assert resp[1][1:] == ["9-9-9", "7", ""]                      # block 2 (2 intervals)

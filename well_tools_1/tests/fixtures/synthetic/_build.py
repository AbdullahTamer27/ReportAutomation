"""Build the synthetic fixture set (deterministic).

Run this once to (re)generate the committed inputs in this folder:

    python tests/fixtures/synthetic/_build.py

Produces:
    template.docx   — a tiny universal template exercising the no-config path:
                      text tags, a company-conditional line, a {{DISC}} table,
                      a repeatable damage block, and a well-head overlay box.
    data.xlsx       — an empty workbook (the template has no table tags, so the
                      table pass is a no-op — enough to drive the pipeline).
    inputs.json     — the non-file inputs passed to report_service.generate.

The golden test generates a report from these and diffs it against the committed
`golden.docx`. If you change this builder, regenerate the golden (the test tells
you how).
"""

import json
import os

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import openpyxl

HERE = os.path.dirname(os.path.abspath(__file__))

_NS = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
       'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
       'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
       'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
       'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"')


def _textbox_p(tag):
    """A paragraph holding a DrawingML text box whose visible text is `tag`."""
    return (
        f'<w:p {_NS}><w:r><mc:AlternateContent><mc:Choice Requires="wps"><w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="1" '
        f'behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/><wp:positionH relativeFrom="column"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="paragraph"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="1000000" cy="300000"/><wp:effectExtent l="0" t="0" r="0" b="0"/><wp:wrapNone/>'
        f'<wp:docPr id="1" name="tb"/><a:graphic>'
        f'<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<wps:wsp><wps:cNvSpPr txBox="1"/><wps:spPr/><wps:txbx><w:txbxContent>'
        f'<w:p><w:r><w:t>{tag}</w:t></w:r></w:p></w:txbxContent></wps:txbx><wps:bodyPr/></wps:wsp>'
        f'</a:graphicData></a:graphic></wp:anchor></w:drawing></mc:Choice></mc:AlternateContent></w:r></w:p>'
    )


def build():
    # --- template.docx ---
    doc = Document()
    doc.add_paragraph("Well: {{well_name}}   Type: {{well_type}}   Field: {{field}}")
    doc.add_paragraph("Bottom depth: {{btm_depth}}   Delivery: {{delivery_date}}")
    doc.add_paragraph("Logged {{log_date}} · orig {{orig_comp}} · last WKO {{last_wko}}")
    doc.add_paragraph("KEEP-CORR {{weatherford_corr}} END")   # conditional line

    body = doc.element.body
    sect = body.find(qn("w:sectPr"))

    def addp(text):
        sect.addprevious(parse_xml(
            f'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f'<w:r><w:t>{text}</w:t></w:r></w:p>'))

    # Well-head overlay box (filled by the overlays pass).
    sect.addprevious(parse_xml(_textbox_p("{{ovl_wellhead}}")))
    # Repeatable damage block with an image placeholder tag.
    addp("{{damage_block_start}}")
    addp("Damage photo: {{DMG@N_1}}")
    addp("{{damage_block_end}}")

    # A {{DISC}} disclaimer table (removed when include_disclaimer is False).
    tbl = doc.add_table(rows=1, cols=1)
    tbl.rows[0].cells[0].text = "{{DISC}} Standard disclaimer text."

    doc.save(os.path.join(HERE, "template.docx"))

    # --- data.xlsx (empty; no table tags to fill) ---
    openpyxl.Workbook().save(os.path.join(HERE, "data.xlsx"))

    # --- inputs.json ---
    inputs = {
        "config": None,
        "company_name": "ACME",
        "company_logo": None,
        "damage_count": 1,
        "wellhead_damage": True,
        "include_disclaimer": False,
        "well_name": "SYNTH_1_0",
        "well_type": "Oil producer",
        "btm_depth": "7000 ft",
        "field": "Testland",
        "log_date": "2020-09-09",
        "orig_comp": "1981/07/13",
        "last_wko": "2013/09/06 #4",
        "xml": None,
    }
    with open(os.path.join(HERE, "inputs.json"), "w") as f:
        json.dump(inputs, f, indent=2)

    print("Synthetic fixture written to", HERE)


if __name__ == "__main__":
    build()
